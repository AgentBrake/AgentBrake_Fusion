from __future__ import annotations

from pathlib import Path
import math

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent / "generated"
CHART_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "AgentBrake-Fusion_experiment_chapter.docx"


SOURCE_FILES = [
    "AgentBrake_PPT_Experiment_Slides_Content_Tables_with_Latency_Ablation.docx",
    "RepoShield_AgentDojo_Full_E2E_Experiment.docx",
    "RepoShield_AgentDojo_QwenPlus_Full_E2E_Experiment.docx",
    "RepoShield实验一报告.docx",
]


def pct(x: float | None) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "N/R"
    return f"{x:.2f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def style_table(table, header_fill: str = "EAF1F8", font_size: float = 8.2) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if ri == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str, font_size: float = 8.0):
    cap = doc.add_paragraph(caption)
    cap.style = doc.styles["CaptionText"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=font_size)
    style_table(table, font_size=font_size)
    doc.add_paragraph()
    return table


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["NoteText"]
    p.add_run("说明：").bold = True
    p.add_run(text)


def make_charts() -> dict[str, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({
        "font.family": "DejaVu Sans",
        "axes.edgecolor": "#334155",
        "axes.labelcolor": "#0f172a",
        "xtick.color": "#0f172a",
        "ytick.color": "#0f172a",
        "figure.dpi": 180,
    })

    # Experiment 1: blocking/pass bar chart.
    fig1, ax1 = plt.subplots(figsize=(5.8, 3.3))
    names = ["Dangerous\nBlocking", "Safe\nPass"]
    vals = [98.55, 93.89]
    colors = ["#0f766e", "#2563eb"]
    bars = ax1.bar(names, vals, color=colors, width=0.55)
    ax1.set_ylim(0, 105)
    ax1.set_ylabel("Rate (%)")
    ax1.set_title("Pre-execution Decision Accuracy")
    ax1.grid(axis="y", color="#e2e8f0", linewidth=0.8)
    for bar, v in zip(bars, vals):
        ax1.text(bar.get_x() + bar.get_width() / 2, v + 1.5, f"{v:.2f}%", ha="center", va="bottom", fontsize=9)
    fig1.tight_layout()
    p1 = CHART_DIR / "exp1_block_pass_bar.png"
    fig1.savefig(p1, bbox_inches="tight")
    plt.close(fig1)

    # Experiment 4: latency line chart.
    latency = [
        ("DS-RS", 0.906, 2.250),
        ("DS-GW", 0.917, 2.132),
        ("DS-OR", 0.919, 2.278),
        ("QW-RS", 1.230, 3.075),
        ("QW-GW", 0.852, 1.206),
        ("QW-OR", 0.985, 3.030),
    ]
    x = np.arange(len(latency))
    fig2, ax2 = plt.subplots(figsize=(7.2, 3.8))
    ax2.plot(x, [r[1] for r in latency], color="#2563eb", marker="o", linewidth=2.2, label="p50 latency")
    ax2.plot(x, [r[2] for r in latency], color="#dc2626", marker="s", linewidth=2.2, label="p95 latency")
    ax2.set_xticks(x)
    ax2.set_xticklabels([r[0] for r in latency])
    ax2.set_ylabel("Latency (ms)")
    ax2.set_title("Pre-execution MSJ Engine Policy Latency")
    ax2.grid(axis="y", color="#e2e8f0", linewidth=0.8)
    ax2.legend(frameon=False, ncols=2, loc="upper left")
    for xi, p50, p95 in [(i, r[1], r[2]) for i, r in enumerate(latency)]:
        ax2.text(xi, p95 + 0.08, f"{p95:.2f}", ha="center", fontsize=8, color="#991b1b")
    fig2.tight_layout()
    p2 = CHART_DIR / "exp4_latency_line.png"
    fig2.savefig(p2, bbox_inches="tight")
    plt.close(fig2)

    # Experiment 6: suite-level mixed chart using only source-reported by-suite fields.
    suites = ["banking", "travel", "workspace", "slack"]
    deepseek = {
        "banking": {"asr": 0.00, "strict": 84.03, "gateway": 79.17, "oracle": 81.25},
        "travel": {"asr": 0.00, "strict": 72.86, "gateway": 72.14, "oracle": 77.14},
        "workspace": {"asr": 0.18, "strict": 81.79, "gateway": 81.43, "oracle": 81.25},
        "slack": {"asr": 0.95, "strict": 31.43, "gateway": 30.48, "oracle": 30.48},
    }
    qwen = {
        "banking": {"asr": 0.00, "strict": 56.25, "gateway": 58.33, "oracle": 53.47},
        "travel": {"asr": 0.00, "strict": 27.14, "gateway": 25.71, "oracle": 27.86},
        "workspace": {"asr": 0.36, "strict": 78.39, "gateway": 76.79, "oracle": 78.57},
        "slack": {"asr": 0.00, "strict": 21.90, "gateway": 20.95, "oracle": 25.71},
    }
    fig3, axes = plt.subplots(2, 2, figsize=(8.0, 5.8), sharey=False)
    methods = ["strict", "gateway", "oracle"]
    labels = ["Strict", "Gateway", "Oracle"]
    legend_handles = None
    legend_labels = None
    for ax, suite in zip(axes.flatten(), suites):
        width = 0.36
        idx = np.arange(len(methods))
        ds_utils = [deepseek[suite][m] for m in methods]
        qw_utils = [qwen[suite][m] for m in methods]
        ax2 = ax.twinx()
        ax.bar(idx - width / 2, ds_utils, width, color="#93c5fd", label="DeepSeek utility")
        ax.bar(idx + width / 2, qw_utils, width, color="#fca5a5", label="Qwen utility")
        l1, = ax2.plot(idx, [deepseek[suite]["asr"]] * len(methods), color="#1d4ed8", marker="o", linewidth=1.8, label="DeepSeek strict ASR")
        l2, = ax2.plot(idx, [qwen[suite]["asr"]] * len(methods), color="#b91c1c", marker="s", linewidth=1.8, label="Qwen strict ASR")
        ax.set_title(suite)
        ax.set_xticks(idx)
        ax.set_xticklabels(labels, rotation=15)
        ax.set_ylim(0, 100)
        ax2.set_ylim(0, 5)
        ax2.set_yticks([0, 1, 2, 3, 4, 5])
        ax.grid(axis="y", color="#e2e8f0", linewidth=0.7)
        if suite in {"banking", "workspace"}:
            ax.set_ylabel("Reported Utility (%)")
        if suite in {"travel", "slack"}:
            ax2.set_ylabel("Strict ASR (%)")
        if legend_handles is None:
            handles, labels_ = ax.get_legend_handles_labels()
            legend_handles = handles + [l1, l2]
            legend_labels = labels_ + ["DeepSeek strict ASR", "Qwen strict ASR"]
    fig3.legend(legend_handles, legend_labels, loc="lower center", ncols=4, frameon=False)
    fig3.suptitle("Suite-level Behavior (source-reported utility; strict ASR)", y=0.98)
    fig3.tight_layout(rect=(0, 0.08, 1, 0.95))
    p3 = CHART_DIR / "exp6_suite_mixed.png"
    fig3.savefig(p3, bbox_inches="tight")
    plt.close(fig3)

    return {"exp1": p1, "latency": p2, "suite": p3}


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12.5, "1F4E79"),
        ("Heading 3", 11, "334155"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    if "CaptionText" not in styles:
        styles.add_style("CaptionText", 1)
    cap = styles["CaptionText"]
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cap.font.size = Pt(9)
    cap.font.bold = True
    cap.font.color.rgb = RGBColor.from_string("334155")
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    if "NoteText" not in styles:
        styles.add_style("NoteText", 1)
    note = styles["NoteText"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor.from_string("475569")
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(6)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AgentBrake-Fusion 实验评估")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("顶会论文风格实验章节草稿：安全性、可用性、恢复能力、延迟与消融分析")
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor.from_string("475569")
    doc.add_paragraph()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    chart_paths = make_charts()
    doc = Document()
    configure_doc(doc)
    add_title(doc)

    doc.add_heading("0. 数据来源与写作约束", level=1)
    doc.add_paragraph(
        "本实验章节只使用用户提供的 Word 数据文件中的已报告数值。若某项指标在源文件中未出现，本文以 N/R "
        "（not reported）标注，并在对应实验段落中说明未进行估算或外推。"
    )
    add_table(doc, ["源文件", "在本文中的用途"], [[f, "提取实验设置、结果表、恢复/确认、延迟或消融数据"] for f in SOURCE_FILES], "表 0. 用户提供的数据文件", 8.0)

    doc.add_heading("1. 实验设置与研究问题", level=1)
    doc.add_paragraph(
        "我们围绕六个问题评估 AgentBrake-Fusion：执行前是否能区分危险动作和安全动作；在 AgentDojo "
        "全量任务中是否兼顾安全性与可用性；阻断或确认之后 agent 是否仍能恢复合法任务；执行前 MSJ Engine "
        "延迟是否足够低；MSJ Engine 与 ActionGraph 的关键组件分别贡献什么；以及四个 AgentDojo 场景中防御表现是否一致。"
    )
    add_table(
        doc,
        ["RQ", "问题", "对应实验"],
        [
            ["RQ1", "执行前裁决能否阻断危险动作并放行安全动作？", "实验一"],
            ["RQ2", "AgentDojo 全量任务中不同防御的安全-可用性权衡如何？", "实验二"],
            ["RQ3", "block / require_confirmation 后合法任务是否可恢复？", "实验三"],
            ["RQ4", "MSJ Engine 的执行前延迟是否满足网关要求？", "实验四"],
            ["RQ5", "MSJ Engine 与 ActionGraph 组件分别贡献什么？", "实验五"],
            ["RQ6", "banking / travel / workspace / slack 场景中的差异是什么？", "实验六"],
        ],
        "表 1. 研究问题与实验映射",
        8.0,
    )

    doc.add_heading("2. 实验一：危险动作拦截率与安全动作放行率", level=1)
    doc.add_paragraph(
        "该实验以 trace-derived replay cases 为样本，验证 AgentBrake-Fusion 在工具执行前是否能区分危险动作和安全动作。"
        "源文件报告共 200 个正式 replay cases，其中 unsafe cases 为 69，safe cases 为 131，所有正式样本均能追溯到原始 full trace。"
    )
    add_table(
        doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["Candidate traces considered", "307", "候选轨迹规模"],
            ["Formal replay cases", "200", "正式计入的 replay cases"],
            ["Unsafe cases", "69", "危险工具调用样本"],
            ["Safe cases", "131", "安全工具调用样本"],
            ["Dangerous Action Blocking Rate", "98.55%", "危险工具调用被 block 的比例"],
            ["Safe Action Pass Rate", "93.89%", "安全工具调用被 allow 的比例"],
            ["False Positive Rate", "5.34%", "安全工具调用被 hard block 的比例"],
            ["Block Reason Accuracy", "100.00%", "阻断 reason code 与风险类型匹配比例"],
        ],
        "表 2. 实验一核心结果",
        8.0,
    )
    p = doc.add_paragraph("图 1. Dangerous Action Blocking Rate 与 Safe Action Pass Rate")
    p.style = doc.styles["CaptionText"]
    doc.add_picture(str(chart_paths["exp1"]), width=Inches(4.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "结果表明，AgentBrake-Fusion 在工具执行前对危险动作具有较高拦截能力，同时对正常授权工具调用保持较高放行率。"
        "错误案例集中在 workspace 的 send_email 安全样本误拦，以及 1 个 workspace unsafe send_email 被 allow。"
    )

    doc.add_heading("3. 实验二：AgentDojo 全量实验", level=1)
    doc.add_paragraph(
        "该实验在 AgentDojo 全量任务上比较 no_defense、tool_filter、MELON、Progent、DRIFT 与 AgentBrake-Fusion 三种运行模式。"
        "指标包括 ASR、Security、User Utility 与 Secure Utility。表中外部方法 MELON/Progent/DRIFT 来自用户提供的对比表；"
        "no_defense、tool_filter 与 AgentBrake-Fusion 来自用户提供的 DeepSeek/Qwen 全量 E2E 实验报告。"
    )
    exp2_rows = [
        ["deepseek-v4-flash", "no_defense", "16.12%", "83.88%", "68.60%", "63.01%"],
        ["deepseek-v4-flash", "tool_filter", "0.00%", "100.00%", "5.06%", "5.06%"],
        ["deepseek-v4-flash", "MELON", "0.84%", "99.16%", "34.25%", "33.40%"],
        ["deepseek-v4-flash", "Progent", "9.00%", "91.00%", "73.68%", "69.78%"],
        ["deepseek-v4-flash", "DRIFT", "0.74%", "99.26%", "64.81%", "64.59%"],
        ["deepseek-v4-flash", "AgentBrake_strict", "0.21%", "99.79%", "75.24%", "75.24%"],
        ["deepseek-v4-flash", "AgentBrake_gateway_eval", "0.00%", "100.00%", "74.08%", "74.08%"],
        ["deepseek-v4-flash", "AgentBrake_oracle_user_eval", "0.00%", "100.00%", "75.03%", "75.03%"],
        ["qwen-plus", "no_defense", "29.19%", "70.81%", "53.32%", "43.31%"],
        ["qwen-plus", "tool_filter", "0.00%", "100.00%", "4.74%", "4.74%"],
        ["qwen-plus", "MELON", "0.42%", "99.58%", "26.13%", "25.71%"],
        ["qwen-plus", "Progent", "1.18%", "98.82%", "72.07%", "70.89%"],
        ["qwen-plus", "DRIFT", "3.58%", "96.42%", "75.34%", "73.87%"],
        ["qwen-plus", "AgentBrake_strict", "0.21%", "99.79%", "61.22%", "61.22%"],
        ["qwen-plus", "AgentBrake_gateway_eval", "0.32%", "99.68%", "60.27%", "60.27%"],
        ["qwen-plus", "AgentBrake_oracle_user_eval", "0.32%", "99.68%", "61.43%", "61.43%"],
    ]
    add_table(doc, ["Model", "Defense", "ASR ↓", "Security ↑", "User Utility ↑", "Secure Utility ↑"], exp2_rows, "表 3. AgentDojo 全量任务主结果", 7.4)
    doc.add_paragraph(
        "从表 3 可见，tool_filter 能把 ASR 降为 0，但 User Utility 仅为 5.06% / 4.74%，说明简单禁用工具会严重破坏任务完成。"
        "AgentBrake_strict 在两个模型上的 ASR 均为 0.21%，同时保持比 tool_filter 高得多的 Secure Utility。"
        "与外部方法相比，AgentBrake-Fusion 在 DeepSeek 上兼具最低 ASR 之一与最高 Secure Utility；在 Qwen 上则以较低 ASR 保持超过 60% 的 Secure Utility。"
    )

    doc.add_heading("4. 实验三：恢复与确认实验", level=1)
    doc.add_paragraph(
        "该实验检验安全拦截或确认之后，agent 是否仍能继续完成合法用户任务。源文件报告了 blocked rate、confirm rate、post-block success、"
        "recovery success、confirm execute 和 repeated block 等指标。源文件未单独报告名为 post_block_secure_utility 的指标；本文保留源文件中的 "
        "Post-block Success，并将 post_block_secure_utility 标为 N/R。"
    )
    rec_rows = [
        ["deepseek-v4-flash", "AgentBrake_strict", "12.54%", "1.58%", "6.72%", "6.72%", "0.00%", "0.00%", "N/R"],
        ["deepseek-v4-flash", "AgentBrake_gateway_eval", "10.85%", "5.58%", "0.97%", "15.07%", "100.00%", "2.91%", "N/R"],
        ["deepseek-v4-flash", "AgentBrake_oracle_user_eval", "11.38%", "5.58%", "2.78%", "15.65%", "100.00%", "2.78%", "N/R"],
        ["qwen-plus", "AgentBrake_strict", "11.06%", "0.21%", "32.38%", "32.38%", "0.00%", "0.00%", "N/R"],
        ["qwen-plus", "AgentBrake_gateway_eval", "10.22%", "2.95%", "27.84%", "31.71%", "100.00%", "0.00%", "N/R"],
        ["qwen-plus", "AgentBrake_oracle_user_eval", "10.54%", "2.95%", "33.00%", "35.71%", "100.00%", "0.00%", "N/R"],
    ]
    add_table(
        doc,
        ["Model", "Mode", "blocked", "confirm_required", "post-block success", "recovery success", "confirm_execute", "repeated block", "post_block_secure_utility"],
        rec_rows,
        "表 4. 恢复与确认结果",
        6.6,
    )
    doc.add_paragraph(
        "DeepSeek 上，gateway_eval 与 oracle_user_eval 的恢复成功率高于 strict，说明确认路径能够为部分被阻断任务提供继续执行机会，但 repeated block 约为 2.8%--2.9%，仍存在反复阻断。"
        "Qwen 上 repeated block 为 0，且 oracle_user_eval 的 recovery success 达到 35.71%，说明理想确认能提升可用性上界。"
        "不过确认执行率为 100% 并不意味着任务必然恢复，表明恢复质量还受模型规划能力和后续工具链影响。"
    )

    doc.add_heading("5. 实验四：执行前 MSJ Engine 延迟实验", level=1)
    doc.add_paragraph(
        "该实验验证多源证据裁决是否足够轻量，能否作为执行前安全网关使用。源文件仅报告 policy_p50_ms 与 policy_p95_ms，"
        "并明确该指标不包含 LLM API 请求、外部工具执行、网络往返和 trace 写盘。average latency、p99 latency、"
        "以及按 evidence items / tool calls / graph facts 分桶的延迟未在源文件中报告，因此本文不生成对应数值。"
    )
    latency_rows = [
        ["deepseek-v4-flash", "AgentBrake_strict", "949", "N/R", "0.906 ms", "2.250 ms", "N/R"],
        ["deepseek-v4-flash", "AgentBrake_gateway_eval", "949", "N/R", "0.917 ms", "2.132 ms", "N/R"],
        ["deepseek-v4-flash", "AgentBrake_oracle_user_eval", "949", "N/R", "0.919 ms", "2.278 ms", "N/R"],
        ["qwen-plus", "AgentBrake_strict", "949", "N/R", "1.230 ms", "3.075 ms", "N/R"],
        ["qwen-plus", "AgentBrake_gateway_eval", "949", "N/R", "0.852 ms", "1.206 ms", "N/R"],
        ["qwen-plus", "AgentBrake_oracle_user_eval", "949", "N/R", "0.985 ms", "3.030 ms", "N/R"],
    ]
    add_table(doc, ["Model", "Mode", "N", "Average", "p50", "p95", "p99"], latency_rows, "表 5. MSJ Engine 执行前 policy 延迟", 7.2)
    p = doc.add_paragraph("图 2. MSJ Engine policy p50 / p95 延迟")
    p.style = doc.styles["CaptionText"]
    doc.add_picture(str(chart_paths["latency"]), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "所有已报告 p50 均在 1.23 ms 以内，p95 最高为 3.075 ms。相对于一次 LLM API 调用或外部工具调用，这一量级支持将 MSJ Engine 放在工具执行前作为同步安全裁决网关。"
    )

    doc.add_heading("6. 实验五：消融实验", level=1)
    doc.add_paragraph(
        "消融实验使用 500 个 AgentDojo 诊断样本。源文件说明这些样本来自 AgentDojo E2E 轨迹，并覆盖 attack_active、blocked_critical 与 safe_side_effect_control。"
        "样本能够覆盖参数污染、私密数据外传、跨步影响、恢复确认等风险，是因为其正式样本中同时包含金融提交、外部通信、文件/日历副作用、成员扩展、私密数据外传和正常用户授权动作。"
        "源文件报告的 violation_type 包括 private_data_exfiltration、private_data_exfiltration_to_external_sink、financial_commit_from_untrusted_source、"
        "membership_expansion_from_untrusted_source、calendar_side_effect、destructive_file_operation 等。"
    )
    ablation_rows = [
        ["full", "0.40%", "99.28%", "58.00%", "58.00%", "32.38%", "21.00%", "完整 MSJ Engine + ActionGraph 基线"],
        ["rule_only", "0.20%", "99.64%", "16.20%", "16.20%", "7.94%", "68.00%", "接近零 ASR 主要来自过度阻断，utility 大幅下降"],
        ["no_binding", "0.00%", "100.00%", "17.00%", "17.00%", "8.59%", "65.20%", "去掉任务/参数绑定后误拦增多"],
        ["no_recovery_guidance", "0.40%", "99.28%", "57.60%", "57.60%", "24.47%", "18.80%", "主要降低阻断后恢复成功"],
        ["flatten_action_graph", "0.80%", "98.56%", "57.00%", "57.00%", "28.74%", "44 flips", "图结构缺失导致 ASR 上升"],
        ["no_actiongraph_provenance_edges", "0.40%", "99.28%", "58.60%", "58.60%", "28.43%", "45 flips", "来源追踪影响裁决边界"],
        ["no_actiongraph_dataflow_edges", "0.60%", "98.92%", "59.40%", "59.40%", "32.58%", "44 flips", "数据流边影响外传识别"],
        ["no_actiongraph_history_edges", "0.80%", "98.56%", "59.60%", "59.60%", "30.77%", "49 flips", "历史边影响恢复稳定性"],
    ]
    add_table(
        doc,
        ["Variant", "ASR ↓", "Suppression ↑", "User Utility ↑", "Secure Utility ↑", "Post-block ↑", "Block/Flip", "Conclusion"],
        ablation_rows,
        "表 6. MSJ Engine 与 ActionGraph 消融结果",
        6.3,
    )
    doc.add_paragraph(
        "消融结果显示，rule_only 和 no_binding 虽然能维持极低 ASR，但会把 Secure Utility 压低到 16%--17%，说明多源裁决不能退化为单纯规则阻断。"
        "no_recovery_guidance 对总体 utility 影响较小，但 Post-block Success 从 32.38% 降至 24.47%，说明恢复引导的贡献集中在阻断后的继续执行。"
        "ActionGraph 相关变体的 decision flips 均在 44--49 之间，其中 flatten_action_graph 与 no_actiongraph_history_edges 将 ASR 推高到 0.80%，支持图结构和历史边对安全稳定性的贡献。"
    )

    doc.add_page_break()
    doc.add_heading("7. 实验六：四个 AgentDojo 场景分场景分析", level=1)
    doc.add_paragraph(
        "源文件按 suite 报告了 no-defense ASR、strict ASR，以及 strict/gateway/oracle 三种模式的 utility。"
        "源文件未分别给出 gateway/oracle 的分场景 ASR，也未把分场景 utility 拆成 User Utility 与 Secure Utility；"
        "因此图 3 使用已报告的 strict ASR 作为场景攻击残余指标，并以三种模式的 source-reported utility 展示可用性差异。"
    )
    suite_rows = [
        ["deepseek-v4-flash", "banking", "144", "13.19%", "0.00%", "84.03%", "79.17%", "81.25%"],
        ["deepseek-v4-flash", "travel", "140", "12.14%", "0.00%", "72.86%", "72.14%", "77.14%"],
        ["deepseek-v4-flash", "workspace", "560", "13.93%", "0.18%", "81.79%", "81.43%", "81.25%"],
        ["deepseek-v4-flash", "slack", "105", "37.14%", "0.95%", "31.43%", "30.48%", "30.48%"],
        ["qwen-plus", "banking", "144", "52.08%", "0.00%", "56.25%", "58.33%", "53.47%"],
        ["qwen-plus", "travel", "140", "30.71%", "0.00%", "27.14%", "25.71%", "27.86%"],
        ["qwen-plus", "workspace", "560", "13.57%", "0.36%", "78.39%", "76.79%", "78.57%"],
        ["qwen-plus", "slack", "105", "79.05%", "0.00%", "21.90%", "20.95%", "25.71%"],
    ]
    add_table(
        doc,
        ["Model", "Suite", "Cases", "No-defense ASR", "Strict ASR", "Strict Utility", "Gateway Utility", "Oracle Utility"],
        suite_rows,
        "表 7. 四场景分场景结果",
        6.8,
    )
    p = doc.add_paragraph("图 3. AgentDojo 四场景组合图：utility bars + strict ASR line")
    p.style = doc.styles["CaptionText"]
    doc.add_picture(str(chart_paths["suite"]), width=Inches(6.45))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "banking 与 workspace 的 utility 明显高于 slack/travel，说明前两类任务的授权目标和工具参数更容易被上下文证据绑定。"
        "slack 在两个模型上 utility 都较低，原因可能是私信、频道消息、成员变更等外部通信动作更容易触发保守裁决；"
        "travel 在 Qwen 上 utility 偏低，说明预订类任务对模型规划与阻断后恢复较敏感。"
    )

    doc.add_heading("8. 结论摘要", level=1)
    doc.add_paragraph(
        "综合六组实验，AgentBrake-Fusion 在执行前工具边界上能够高精度拦截危险动作并保留安全动作，"
        "在 AgentDojo 全量任务中显著降低 ASR，同时比简单 tool_filter 保留更高任务可用性。恢复与确认实验表明 require_confirmation "
        "能够改善部分被阻断任务的继续执行，但恢复效果仍受模型和场景影响。延迟结果显示已报告 p50/p95 均处于毫秒级，"
        "适合作为执行前安全网关。消融结果进一步说明，MSJ Engine 的 task binding、recovery guidance 以及 ActionGraph 的结构化边共同支撑了安全-可用性平衡。"
    )

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOCX)
